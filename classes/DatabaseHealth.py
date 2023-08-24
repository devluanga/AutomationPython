import smtplib
import datetime
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

import sqlalchemy as sa
import pandas as pd
from docx.shared import Pt, RGBColor
import requests
from requests.auth import HTTPBasicAuth
import json

class DatabaseHealth:
    cx_Oracle = None
    def __init__(self, cx_Oracle):
        self.cx_Oracle = cx_Oracle
    def db_connection(self, username, password, hostname, service_name, port):
        try:
            oracle_connection_string_fmt = (
                    'oracle+cx_oracle://{username}:{password}@' +
                    self.cx_Oracle.makedsn('{hostname}', '{port}', service_name='{service_name}')
            )
            url = oracle_connection_string_fmt.format(
                username=username, password=password,
                hostname=hostname, port=port,
                service_name=service_name,
            )
            engine: sa.engine.Engine = sa.create_engine(url, echo=False, arraysize=1000)
            # logging.info('Successfully connected to ' + service_name + ' on host ' + hostname)
        except(ConnectionError):
            print("Error")
            # logging.error('Failed to connect to ' + service_name + ' on host ' + hostname)
        return engine

    def doc_insert_table(self,doc, query,  engine=None, fontsize=5, tableStyle="Colorful Grid Accent 5"):
        data = pd.read_sql(query, engine)
        columns = (data.columns).tolist()
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(query, 'Emphasis')
        font = run.font
        font.bold = True
        font.size = Pt(fontsize)

        table = doc.add_table(data.shape[0] + 1, data.shape[1])
        col = table.rows[0].cells
        for i in range(0, len(columns)):
            col[i].text = (str(columns.__getitem__(i))).upper()
        for i in range(data.shape[0]):
            for j in range(data.shape[-1]):
                table.cell(i + 1, j).text = str(data.values[i, j])
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                paragraph = paragraphs[0]
                run_obj = paragraph.runs
                run = run_obj[0]
                font = run.font
                font.size = Pt(fontsize)
        table.style = tableStyle


    def checkurlstatus(self, doc, url):
        try:
            requests.packages.urllib3.disable_warnings()
            r = requests.head(url, verify=False)
            if r.status_code == 200:
                doc.add_heading(url+" OK", level=8)
            else:
                doc.add_heading(url+" error"+r.status_code, level=8)
        except requests.ConnectionError:
            doc.add_heading(url + " Failed", level=8)

    def api_call(self, doc, url, username, password, requestBody):
        requests.packages.urllib3.disable_warnings()
        table = doc.add_table(rows=3, cols=2, style="Colorful Shading Accent 6")

        try:
            r = requests.post(url, auth=HTTPBasicAuth(username, password), verify=False,json=json.loads(requestBody))
            # Adding heading in the 1st row of the table
            f_row = table.rows[0].cells
            fc = table.cell(0, 0)
            sc = table.cell(0, 1)
            hr = fc.merge(sc)

            f_row[0].text = ""
            run = f_row[0].paragraphs[0].add_run("POST:" + url)
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
            s_row = table.rows[1].cells
            s_row[0].text = "Request Body"
            s_row[1].text = "Response Body"

            t_row = table.rows[2].cells
            run = t_row[0].paragraphs[0].add_run(requestBody)
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(144, 238, 144)
            t_row[1].text = r.text
            # row1[1].text = r.text
            print(r.json())
        except Exception as e:
            f_row = table.rows[0].cells
            fc = table.cell(0, 0)
            sc = table.cell(0, 1)
            hr = fc.merge(sc)

            f_row[0].text = ""
            run = f_row[0].paragraphs[0].add_run("POST:" + url)
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
            s_row = table.rows[1].cells
            s_row[0].text = "Request Body"
            s_row[1].text = "Response Body"

            t_row = table.rows[2].cells
            run = t_row[0].paragraphs[0].add_run(requestBody)
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(144, 238, 144)
            t_row = table.rows[2].cells
            t_row[1].text = str(e)

    def send_message(self, filename, subject, sender, recipients, body, username, password):

        msg = MIMEMultipart()

        msg["Subject"] = subject
        msg["From"] = sender
        msg["To"] = ", ".join(recipients)
        body+=" Generated on "+str(datetime.date.today()) + " @ "+str(datetime.datetime.now())
        msg.attach(MIMEText(body, "plain"))

        msg.add_header("Content-Type", "text/html")
        msg.as_string()

        with open(filename, "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f"attachment; filename= {filename.replace('documents', '')}",
            )
        msg.attach(part)
        try:
            mail_server = smtplib.SMTP('webmail.safaricom.co.ke', 587)
            mail_server.ehlo()
            mail_server.starttls()
            mail_server.ehlo()
            mail_server.login(username, password)
            mail_server.send_message(msg)
            mail_server.quit()
        except smtplib.SMTPException as exception:
            print(exception)
        except smtplib.SMTPAuthenticationError as auth_error:
            print(auth_error)
        except smtplib.SMTPConnectError as conn_error:
            print(conn_error)
