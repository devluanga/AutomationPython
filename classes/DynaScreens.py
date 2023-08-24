from selenium.common.exceptions import TimeoutException, NoSuchElementException
from selenium.webdriver.common.by import By
import time
from selenium import webdriver
from selenium.webdriver.firefox.options import Options
import os
from datetime import date
import traceback
from docx.shared import Pt, Inches
class DynatraceScreenshots:
    browser = None
    options = Options()
    options.headless = True
    image_path = ""
    def __init__(self, driver_path, image_path):
        self.browser = webdriver.Firefox(options=self.options, executable_path=driver_path)
        self.image_path=image_path
    def checkFolderExists(self):
        try:
            os.makedirs(self.image_path, exist_ok=True)
            print("--confirming "+self.image_path+ " Exists")
            return True
        except Exception:
            print("--could not create directory")
    def dynatraceLogin(self, url, uname):
        print("--Navigating to "+url)
        x = 0;
        while x < 3:
            try:
                self.browser.get(url)
                print("--Attempting to login " + str(x+1))
                username = self.browser.find_element(By.ID, "email_verify")
                nextBtn = self.browser.find_element(By.ID, "next_button")
                username.send_keys(uname)
                nextBtn.click()
                time.sleep(20)
                print("click window")
                self.browser.switch_to.alert.accept()
                print("clicked window")

                x = 5
                print("--logged in to dynatrace")
            except Exception:
                x = x + 1
                traceback.print_exc()
                print("logged in failed trying " + str(x))
                return 0
    def ScreenShotElement(self, elementName, elementtag, elementLocation, mainDashboard, waitTime):
        try:
            if(mainDashboard != ""):
                self.browser.get(mainDashboard)
                print("Loading Dashboard for "+waitTime+" secs")
                time.sleep(int(waitTime))
                print("---Generating " + elementName)
                if(elementtag=="xpath"):
                     element = self.browser.find_element(By.XPATH, elementLocation)
                elif(elementtag=="tagname"):
                    element = self.browser.find_element(By.TAG_NAME, elementLocation)
                else:
                    print("Tagname not defined")
                element.screenshot(self.image_path + os.sep + elementName + ".png")
                print("---Finished Generating " + elementName)
            else:
                print("---Generating " + elementName)
                if (elementtag == "xpath"):
                    element = self.browser.find_element(By.XPATH, elementLocation)
                elif (elementtag == "tagname"):
                    element = self.browser.find_element(By.TAG_NAME, elementLocation)
                else:
                    print("Tagname not defined")
                element.screenshot(self.image_path + os.sep + elementName + ".png")
                print("---Finished Generating " + elementName)
        except NoSuchElementException:
            print("Failed to find "+elementName+ " Skipping ...")
            return False
        except TimeoutException:
            print("Time out")

    def doc_insert_image(self, doc, t_rows, t_cols, images, newPage):
        table = doc.add_table(rows=t_rows, cols=t_cols)
        row_cells = table.add_row().cells
        for im in images:
            run = table.cell(int(im["img_row"]), int(im["img_col"])).paragraphs[0].add_run()
            picture = run.add_picture(im["img_name"])
            picture.height = Inches(float(im["img_height"]))
            picture.width = Inches(float(im["img_width"]))
