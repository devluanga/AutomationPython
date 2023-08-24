from classes import DatabaseHealth, DynaScreens
import configparser
import xml.etree.ElementTree as ET
import cx_Oracle
import logging
import docx
import os
import datetime
from docx.shared import Cm
import time

from cryptography.fernet import Fernet

sep_ = os.sep
config = configparser.ConfigParser()
config.read('configs'+sep_+'config.ini')
tree = ET.parse('xml'+sep_+'screens.xml')
root = tree.getroot()

docxtree = ET.parse('xml'+sep_+'docx.xml')
docxRoot = docxtree.getroot()

dynatraceUrl = config.get('DYNATRACEURL', 'url')
dynatraceUser = config.get('DYNATRACEUSER', 'username')

cx_Oracle.init_oracle_client(lib_dir="drivers"+sep_+"oracle"+sep_+"instantclient_19_9")
logging.basicConfig(filename='healthcheck.log', level=logging.DEBUG)
doc = docx.Document()


today = datetime.date.today()
image_folder = "scrn_" + str(today)
path =  os.getcwd()+os.sep+"screenshots"+os.sep+image_folder

file_name = "healthcheck_"+str(today)+"_"+str(datetime.datetime.now().hour)+".docx"
path_to_driver=os.getcwd()+os.sep+"drivers"+os.sep+"geckodriver.exe"
ds = DynaScreens.DynatraceScreenshots(path_to_driver, path)
dbh = DatabaseHealth.DatabaseHealth(cx_Oracle)

cipher_suite = Fernet(bytes(str(config.get('PROJECT','KEY')),'utf-8'))
# Simswap credentials
sim_swap_username = (cipher_suite.decrypt(bytes(config.get('SIMSWAPDB','sname'), 'utf-8'))).decode()
sim_swap_pass = (cipher_suite.decrypt(bytes(config.get('SIMSWAPDB','secret'), 'utf-8'))).decode()
simswapEngine = dbh.db_connection(sim_swap_username, sim_swap_pass, "svthk1-scan", "SIMSWAP", "1521")
# Subreg credentials
subreg_username = (cipher_suite.decrypt(bytes(config.get('SUBREGDB','sname'), 'utf-8'))).decode()
subreg_pass = (cipher_suite.decrypt(bytes(config.get('SUBREGDB','secret'), 'utf-8'))).decode()
subregEngine=dbh.db_connection(subreg_username, subreg_pass, "svthk2-scan", "SEBREGSB", "1521")
# Eirsb credentials
eirsb_username = (cipher_suite.decrypt(bytes(config.get('EIRSDB','sname'), 'utf-8'))).decode()
eirsb_pass = (cipher_suite.decrypt(bytes(config.get('EIRSDB','secret'), 'utf-8'))).decode()
eirsbEngine =dbh.db_connection(eirsb_username, eirsb_pass, "svthk1-scan", "eirsb", "1521")
# Heko credentials
heko_username = (cipher_suite.decrypt(bytes(config.get('HOKOSTDB','sname'), 'utf-8'))).decode()
heko_pass = (cipher_suite.decrypt(bytes(config.get('HOKOSTDB','secret'), 'utf-8'))).decode()
hekoEngine = dbh.db_connection(heko_username, heko_pass, "svdt5fc1-scan", "HEKODR", "1521")
# Tibco credentials
tibco_username = (cipher_suite.decrypt(bytes(config.get('TIBCODB','sname'), 'utf-8'))).decode()
tibco_pass = (cipher_suite.decrypt(bytes(config.get('TIBCODB','secret'), 'utf-8'))).decode()
tibcodbEngine = dbh.db_connection(tibco_username, tibco_pass, "svdt1-scan2", "tibcodb", "1521")
email_user = config.get('EMAILS','emailusername')
email_pass = (cipher_suite.decrypt(bytes(config.get('EMAILS','emailpassword'), 'utf-8'))).decode()

ds.dynatraceLogin(dynatraceUrl, dynatraceUser)
ds.checkFolderExists()
time.sleep(30)
for dashboard in root:
    mainDashboardUrl = str(dashboard.attrib['url'])
    mainDashboardName = str(dashboard.attrib['name'])
    print("Navigating to " + mainDashboardName)
    for element in dashboard:
        elementName = str(element.attrib["name"])
        waitTime = str(element.attrib["wait"])
        elementTag = str(element.attrib['tag'])
        elementLocation = str(element.attrib['location'])
        ds.ScreenShotElement(elementName, elementTag, elementLocation, mainDashboardUrl, waitTime)
        mainDashboardUrl = ""

for section in docxRoot:
    sectionType = str(section.attrib['type'])
    pageBreak = str(section.attrib["new_page"])
    if (pageBreak == "true"):
        doc.add_page_break()
    if sectionType == "images":
        docImgs = []
        sectionHeading = str(section.attrib['heading'])
        sectionLevel = str(section.attrib['level'])
        sectionRow = str(section.attrib['trow'])
        sectionColumn = str(section.attrib['tcol'])
        doc.add_heading(sectionHeading, level=int(sectionLevel))
        for image in section:
            images = {}
            images["img_name"] = path+os.sep+str(image.attrib["img"])
            images["img_width"] = str(image.attrib['width'])
            images["img_height"] = str(image.attrib['height'])
            images["img_row"] = str(image.attrib['im_row'])
            images["img_col"] = str(image.attrib['im_col'])
            docImgs.append(images)
        print("Appending "+sectionHeading)
        ds.doc_insert_image(doc, t_rows=int(sectionRow), t_cols=int(sectionColumn), images=docImgs, newPage=pageBreak)
    elif sectionType == "database_table":
        tableStyle = "Colorful Grid Accent 5"
        fontsize =5
        if 'style' in section.attrib:
            tableStyle = str((section.attrib['style']))
        if 'fontsize' in section.attrib:
            fontsize = int((section.attrib['fontsize']))

        tableName = str(section.attrib['table'])
        sectionHeading = str(section.attrib['heading'])
        sectionSubHeading = str(section.attrib['subheading'])
        doc.add_heading(sectionHeading, level=2)
        doc.add_paragraph(sectionSubHeading)

        rowLimit = str(section.attrib['limitRows'])
        engine = str(section.attrib['db_engine'])
        subQuery = "SELECT "
        condition=""
        query=""
        sort =""
        for attr in section:
            if attr.tag == "column":
                subQuery+=attr.text+", "
            elif attr.tag == "condition":
                condition=attr.text
            elif attr.tag == "sort":
                sort = attr.text
            elif attr.tag == "query":
                query = attr.text
        if (not query):
            query = "SELECT * FROM ("+subQuery[:-2]+" FROM "+tableName+" "+condition+" "+sort+") WHERE rownum<" +rowLimit
        print(query)
        if engine=="SIMSWAPDB":
            dbh.doc_insert_table(doc, query, simswapEngine,  fontsize, tableStyle)
        elif engine=="SUBREGDB":
            dbh.doc_insert_table(doc, query,subregEngine,  fontsize, tableStyle)
        elif engine=="SIMSWAP":
            dbh.doc_insert_table(doc, query, simswapEngine,  fontsize, tableStyle)
        elif engine=="EIRSDB":
            dbh.doc_insert_table(doc, query, eirsbEngine,  fontsize, tableStyle)
        elif engine=="TIBCODB":
            dbh.doc_insert_table(doc, query, tibcodbEngine,  fontsize, tableStyle)
        elif engine=="HEKODB":
            dbh.doc_insert_table(doc, query, hekoEngine,  fontsize, tableStyle)

    elif sectionType == "api_calls":
        url = ""
        requestBody = ""
        sectionHeading = str(section.attrib['heading'])
        username = 'eai_sms_sender'
        password = 'ea1@SMSsender!'
        for attr in section:
            if attr.tag == "url":
                url=attr.text
            elif attr.tag == "requestbody":
                requestBody=attr.text

        doc.add_heading(sectionHeading, level=2)
        dbh.api_call(doc, url, username, password, requestBody)

#-----------------------------Check URL Status----------------------------------------------------------
logging.debug('BPM AND LB URLS')
doc.add_heading('BPM AND LB URLS STATUS', level=2)

Urls ={
    1: "https://host:port/openspace/?locale=en_US#",
    2: "https://host:port/openspace/?locale=en_US#",
    3: "https://host:port/openspace/?locale=en_US#",
    4: "https://host:port/openspace/?locale=en_US#",
    4: "https://host:port/openspace/?locale=en_US#",
}
for url in Urls.values():
    dbh.checkurlstatus(doc, url)

sections = doc.sections
for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)


doc.save("documents"+os.sep+file_name)
subject=config.get('EMAILS', 'subject')
sender=config.get('EMAILS', 'sender')
recipients=(config.get('EMAILS', 'recipients')).split()
body=config.get('EMAILS', 'body')

dbh.send_message("documents"+os.sep+file_name, subject, sender, recipients, body, email_user, email_pass)



